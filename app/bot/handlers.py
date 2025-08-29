import inspect
import time
import asyncio
import logging
import json
# (same as the "updated_ocr_handlers.zip" version I prepared earlier)
# Included here in full for convenience.
import os, re, time, asyncio, logging, difflib, json
from pathlib import Path
import zipfile
import tempfile
from typing import List, Tuple, Dict, Optional, Any, Union

# --- ensure logging shows our INFO messages ---
_semantic_logger = logging.getLogger("semantic-bot")
if not _semantic_logger.handlers:
    _handler = logging.StreamHandler()
    _formatter = logging.Formatter("[%(levelname)s] %(name)s: %(message)s")
    _handler.setFormatter(_formatter)
    _semantic_logger.addHandler(_handler)
_semantic_logger.setLevel(logging.INFO)
del _handler, _formatter



from telegram import Update, Document, ReplyKeyboardMarkup, InputFile
from telegram.ext import MessageHandler, CommandHandler, filters, ApplicationBuilder

from app.config import (
    TELEGRAM_BOT_TOKEN, DOC_FOLDER, INDEX_FILE, TEXTS_FILE, MANIFEST_FILE,
    DAILY_FREE_LIMIT, TELEGRAM_MSG_LIMIT, RUN_MODE, PUBLIC_BASE_URL, RETRIEVAL_K,
    FORCE_OCR
)
from app.utils.usage import is_admin, get_usage, inc_usage
from app.utils.files import load_manifest, save_manifest, sha256_file
from app.services.indexing import index_file, retrieve_chunks
from app.services.generation import generate_direct_ai_answer, generate_answer_with_gemini
from app.services.tm import tm_process_search, ROW_MATCH_EXPERTISE, ROW_MATCH_REGISTERED, ROW_MATCH_KW
from app.services.extract import extract_text_from_pdf, extract_text_from_docx, extract_text_from_txt, extract_text_from_doc
from app.services.chunking import smart_split_text
from app.services.embeddings import get_embedding
import io
import tempfile
from app.ocr.postprocess import postprocess

from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX

import fitz  # PyMuPDF
from pdf2image import convert_from_path
import pytesseract

logger = logging.getLogger("semantic-bot")

AI_LABEL   = "🤖 AI-чат"
DOCS_LABEL = "FAQ"
WORK_LABEL = "🗂️ Работа с документами"
TM_LABEL   = "🏷️ Товарные знаки"
MAIN_KB    = ReplyKeyboardMarkup([[AI_LABEL, WORK_LABEL, TM_LABEL], [DOCS_LABEL]], resize_keyboard=True)

TM_MODE = "tm"

# ---- OCR defaults tuned for Russian docs ----
OCR_DPI = int(os.getenv("OCR_DPI", "300"))
PDF_COMPRESS = os.getenv("PDF_COMPRESS", "1") == "1"
OCR_MAX_PAGES = int(os.getenv("OCR_MAX_PAGES", "0"))  # 0 = all
POPPLER_PATH = os.getenv("POPPLER_PATH", "").strip() or None
TESS_LANG = os.getenv("TESS_LANG", "rus")  # чисто рус по умолчанию
TESS_CONFIG = os.getenv("TESS_CONFIG", "--oem 3 --psm 6 -c preserve_interword_spaces=1")

# Dedup windows
PDF_DEDUP_TTL = 120
INDEX_DEDUP_TTL = 600

def _split_for_telegram(text: str, max_len: int = TELEGRAM_MSG_LIMIT - 200) -> list[str]:
    parts, buf, cur = [], [], 0
    for p in text.replace("\r\n", "\n").split("\n\n"):
        p = p.strip()
        if not p:
            chunk = "\n\n".join(buf).strip()
            if chunk: parts.append(chunk)
            buf, cur = [], 0
            continue
        need = len(p) + (2 if cur > 0 else 0)
        if cur + need <= max_len:
            buf.append(p); cur += need
        else:
            chunk = "\n\n".join(buf).strip()
            if chunk: parts.append(chunk)
            buf, cur = [], 0
            while len(p) > max_len:
                parts.append(p[:max_len]); p = p[max_len:]
            if p:
                buf, cur = [p], len(p)
    chunk = "\n\n".join(buf).strip()
    if chunk: parts.append(chunk)
    return parts



async def _send_file_safely(update: Update, path: str, download_name: str, caption: str = ""):
    """Send file with size check; auto-zip if too large for Telegram bot (≈50MB limit)."""
    try:
        if not os.path.exists(path):
            await update.message.reply_text("Файл не найден.")
            return
        size = os.path.getsize(path)
        # Telegram ботам лучше держаться < 50 MB
        limit = 45 * 1024 * 1024
        if size > limit:
            # zip on the fly into temp dir
            with tempfile.TemporaryDirectory() as td:
                zip_path = os.path.join(td, download_name + ".zip") if not download_name.endswith(".zip") else os.path.join(td, download_name)
                with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
                    z.write(path, arcname=download_name)
                with open(zip_path, "rb") as f:
                    await update.message.reply_document(document=InputFile(f, filename=os.path.basename(zip_path)),
                                                        caption=(caption + " (упаковано в ZIP)")[:1024])
        else:
            with open(path, "rb") as f:
                await update.message.reply_document(document=InputFile(f, filename=download_name), caption=caption[:1024])
    except Exception as e:
        await update.message.reply_text(f"Ошибка отправки файла: {e}")

async def send_long(update: Update, text: str):
    text = (text or "").strip() or "⚠️ Пустой ответ. Попробуйте переформулировать запрос или загрузить документ заново."
    for c in _split_for_telegram(text):
        await update.message.reply_text(c, disable_web_page_preview=True)

def _normalize_ext(ext: str) -> str:
    mapping = {"с":"c","х":"x","о":"o","р":"p","а":"a","е":"e","к":"k","м":"m","т":"t","н":"h","в":"b"}
    return "".join(mapping.get(ch, ch) for ch in ext)

# ---- Dedup helpers ----
def _pdf_job_key(document: Document) -> str:
    return f"pdf::{document.file_unique_id or document.file_id}"

def _pdf_job_is_running(context, key: str) -> bool:
    jobs = context.user_data.setdefault("pdf_jobs", {})
    rec = jobs.get(key)
    if not rec: return False
    if time.time() - rec.get("ts", 0) > PDF_DEDUP_TTL:
        jobs.pop(key, None)
        return False
    return rec.get("running", False)

def _pdf_job_start(context, key: str, progress_msg_id: int | None):
    jobs = context.user_data.setdefault("pdf_jobs", {})
    jobs[key] = {"running": True, "ts": time.time(), "msg_id": progress_msg_id}

def _pdf_job_finish(context, key: str):
    jobs = context.user_data.setdefault("pdf_jobs", {})
    jobs.pop(key, None)

def _index_job_key(file_unique_id: str | None, file_hash: str | None) -> str:
    return f"idx::{file_unique_id or ''}::{file_hash or ''}"

def _idx_is_running(context, key: str) -> bool:
    jobs = context.bot_data.setdefault("index_jobs", {})
    rec = jobs.get(key)
    if not rec: return False
    if time.time() - rec.get("ts", 0) > INDEX_DEDUP_TTL:
        jobs.pop(key, None)
        return False
    return bool(rec.get("running"))

def _idx_start(context, key: str, msg_id: int | None):
    jobs = context.bot_data.setdefault("index_jobs", {})
    jobs[key] = {"running": True, "ts": time.time(), "msg_id": msg_id}

def _idx_finish(context, key: str):
    jobs = context.bot_data.setdefault("index_jobs", {})
    jobs.pop(key, None)

# ---------- Post-OCR normalization: Latin look-alikes -> Cyrillic ----------
def _normalize_confusables_ru(text: str) -> str:
    """Replace Latin look-alikes with Cyrillic and tidy spaces/punctuation."""
    if not text:
        return text
    mapping = {
        # uppercase
        "A": "А", "B": "В", "C": "С", "E": "Е", "H": "Н", "K": "К", "M": "М",
        "O": "О", "P": "Р", "T": "Т", "X": "Х", "Y": "У",
        # lowercase
        "a": "а", "c": "с", "e": "е", "o": "о", "p": "р", "x": "х", "y": "у",
        # occasional OCR junk
        "Ñ": "Н", "Í": "И", "ì": "и",
    }
    out = []
    for ch in text:
        out.append(mapping.get(ch, ch))
    fixed = "".join(out)
    fixed = re.sub(r"\s+([,.:;!?])", r"\1", fixed)
    fixed = fixed.replace("…", "...")
    return fixed

# ---------- PDF helpers ----------
def _compress_pdf_sync(src: str) -> str:
    try:
        doc = fitz.open(src)
        out = os.path.join(DOC_FOLDER, f"compressed_{Path(src).name}")
        doc.save(out, garbage=4, deflate=True, clean=True, linear=True)
        doc.close()
        return out
    except Exception as e:
        logger.warning("PDF compress failed for %s: %r", src, e)
        return src

def _ocr_pdf_to_text_sync(pdf_path: str, dpi: int, max_pages: int, poppler_path: Optional[str], tess_lang: str) -> str:
    try:
        kwargs = {"dpi": dpi, "fmt": "png"}
        if poppler_path:
            kwargs["poppler_path"] = poppler_path
        images = convert_from_path(pdf_path, **kwargs)
        if max_pages and max_pages > 0:
            images = images[:max_pages]

        texts = []
        for img in images:
            try:
                # Try OpenCV preprocessing if available
                try:
                    import cv2
                    import numpy as np
                    arr = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2GRAY)
                    arr = cv2.threshold(arr, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                    pre_img = arr
                    ocr_text = pytesseract.image_to_string(pre_img, lang=tess_lang, config=TESS_CONFIG)
                except Exception:
                    ocr_text = pytesseract.image_to_string(img, lang=tess_lang, config=TESS_CONFIG)
                ocr_text = _normalize_confusables_ru(ocr_text)
            except Exception as e:
                logger.warning("OCR page failed: %r", e)
                ocr_text = ""
            texts.append(ocr_text)
        return "\n".join(texts).strip()
    except Exception as e:
        logger.error("OCR failed for %s: %r", pdf_path, e)
        return ""

def _save_text_as_docx(text: str, base_name: str) -> str:
    os.makedirs(DOC_FOLDER, exist_ok=True)
    safe = re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", base_name)[:60]
    out = os.path.join(DOC_FOLDER, f"{int(time.time())}_{safe}.docx")
    doc = DocxDocument()
    for line in (text or "").split("\n"):
        doc.add_paragraph(line)
    doc.save(out)
    return out

# ---------- DOCX helpers ----------
def _save_docx_report(filename: str, title: str, body: str) -> str:
    os.makedirs(DOC_FOLDER, exist_ok=True)
    path = os.path.join(DOC_FOLDER, filename)
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for para in (body or "").split("\n"):
        doc.add_paragraph(para)
    doc.save(path)
    return path

def _add_table_change(doc: DocxDocument, old_text: str, new_text: str, caption: str | None = None):
    if caption:
        pcap = doc.add_paragraph()
        pcap.add_run(caption).bold = True
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    # Left: old
    cell_old = table.rows[0].cells[0]
    old_lines = (old_text or '').splitlines() or ['(пусто)']
    for i, line in enumerate(old_lines):
        p = cell_old.paragraphs[0] if i == 0 else cell_old.add_paragraph()
        p.add_run(line)

    # Right: new (yellow)
    cell_new = table.rows[0].cells[1]
    new_lines = (new_text or '').splitlines() or ['(пусто)']
    for i, line in enumerate(new_lines):
        p = cell_new.paragraphs[0] if i == 0 else cell_new.add_paragraph()
        r = p.add_run(line)
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph("")

def _save_docx_compare_tables(filename: str, doc_name_a: str, doc_name_b: str, text_a: str, text_b: str) -> str:
    os.makedirs(DOC_FOLDER, exist_ok=True)
    path = os.path.join(DOC_FOLDER, filename)
    doc = DocxDocument()
    doc.add_heading(f"Изменения: {doc_name_a} → {doc_name_b}", level=1)

    a_lines = [l.rstrip() for l in (text_a or "").splitlines()]
    b_lines = [l.rstrip() for l in (text_b or "").splitlines()]
    sm = difflib.SequenceMatcher(a=a_lines, b=b_lines, autojunk=False)

    change_count = 0
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        change_count += 1
        old_block = "\n".join(a_lines[i1:i2]).strip()
        new_block = "\n".join(b_lines[j1:j2]).strip()
        if tag == "replace":
            _add_table_change(doc, old_block, new_block, caption=f"Различие {change_count}")
        elif tag == "insert":
            _add_table_change(doc, "", new_block, caption=f"Добавлено {change_count}")
        elif tag == "delete":
            _add_table_change(doc, old_block, "", caption=f"Удалено {change_count}")

    if change_count == 0:
        doc.add_paragraph("Изменений не обнаружено.")
    doc.save(path)
    return path

# ---------- Local fallbacks (no LLM) ----------
def _fallback_summary(text: str) -> str:
    text = (text or "").strip()
    if not text:
        return "Документ пуст или не распознан."
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    head = lines[:100]
    def grab(label_variants):
        for lv in label_variants:
            m = re.search(rf"{lv}([^\n\r]{{0,200}})", text, flags=re.IGNORECASE)
            if m:
                return m.group(0).strip()
        return ""
    parties = grab(["стороны", "заказчик", "исполнитель", "арендодатель", "арендатор"])
    term    = grab(["срок", "период действия", "дата окончания"])
    price   = grab(["цена", "стоимость", "вознаграждение", "оплата"])
    subject = grab(["предмет договора", "предмет", "цель договора"])

    bullets = []
    if subject: bullets.append(f"• Предмет: {subject}")
    if parties: bullets.append(f"• Стороны: {parties}")
    if term:    bullets.append(f"• Срок: {term}")
    if price:   bullets.append(f"• Цена/оплата: {price}")
    if not bullets:
        bullets.append("• Ключевые положения не найдены автоматически — возможно, документ нестандартного формата.")

    preview = "\n".join(head[:10])
    return "ИТОГ:\n- Краткое резюме по основным полям.\n\n" + "\n".join(bullets) + (f"\n\nФрагменты:\n{preview}" if preview else "")

def _fallback_check(text: str) -> str:
    text = (text or "").lower()
    if not text:
        return "Документ пуст или не распознан."
    checks = [
        ("Стороны/реквизиты", all(k in text for k in ["реквиз", "подпис", "адрес"]) ),
        ("Предмет договора", "предмет" in text),
        ("Срок и расторжение", any(k in text for k in ["срок", "расторж", "прекращ"])),
        ("Ответственность/неустойка", any(k in text for k in ["ответствен", "неусто", "штраф"])),
        ("Конфиденциальность/персональные данные", any(k in text for k in ["конфиденц", "персональн", "пдн"])),
        ("Форс-мажор", any(k in text for k in ["форс", "непреодол"])),
        ("Порядок оплаты", any(k in text for k in ["оплат", "стоимост", "цена"])),
        ("Права на РИД/рез-ты работ", any(k in text for k in ["исключительн", "право", "интеллектуал"])),
    ]
    lines = [f"• {name}: {'OK' if ok else 'ПРОВЕРИТЬ/ОТСУТСТВУЕТ'}" for name, ok in checks]
    return "ПРОВЕРКА ДОГОВОРА (эвристика без ИИ):\n" + "\n".join(lines)

# ---------- Handlers ----------
async def start(update: Update, context: Any):
    context.user_data["mode"] = "docs"
    usage_left = "∞" if is_admin(update.effective_user.id) else max(0, DAILY_FREE_LIMIT - get_usage(update.effective_user.id))
    msg = (
        "Привет!\n\n"
        "1) 🗂️ Работа с документами — загрузите DOC/DOCX/PDF; для PDF автоматически делаем сжатие+OCR и сохраняем в DOCX.\n"
        "2) Часто задаваемые вопросы — классический режим с глобальным индексом (PDF/DOCX/TXT).\n"
        "3) 🤖 AI-чат — свободный диалог.\n"
        "4) 🏷️ Товарные знаки — поиск по Google Sheets.\n\n"
        f"Сегодняшний лимит AI-чат: {usage_left} сообщений."
    )
    await update.message.reply_text(msg, reply_markup=MAIN_KB)

async def ai_mode(update: Update, context: Any):
    context.user_data["mode"] = "ai"
    usage_left = "∞" if is_admin(update.effective_user.id) else max(0, DAILY_FREE_LIMIT - get_usage(update.effective_user.id))
    await update.message.reply_text(
        f"Режим: AI-чат. Спросите что угодно. Доступно сегодня: {usage_left}.", reply_markup=MAIN_KB
    )

async def docs_mode(update: Update, context: Any):
    context.user_data["mode"] = "docs"
    await update.message.reply_text(
        "Режим: часто задаваемые вопросы по проиндексированным документам. Пришлите файл и задавайте вопрос.", reply_markup=MAIN_KB
    )

async def work_mode(update: Update, context: Any):
    context.user_data["mode"] = "work"
    context.user_data.setdefault("work_docs", [])
    await update.message.reply_text(
        "Режим: 🗂️ Работа с документами.\n\n"
        "Пришлите DOC/DOCX/PDF — добавлю в контекст. Для PDF: сжатие → OCR → сохраняем как DOCX.\n"
        "Команды:\n"
        "• /doc_summary — резюме последнего документа\n"
        "• /doc_check — проверить договор на ошибки/риски\n"
        "• /doc_compare — сравнить два последних загруженных файла (для PDF сравнивается OCR-ДOCX)\n"
        "• /doc_clear — удалить все ранее загруженные документы\n"
        "Любой вопрос в этом режиме — ответ с опорой на загруженные файлы.",
        reply_markup=MAIN_KB
    )

async def tm_mode(update: Update, context: Any):
    context.user_data["mode"] = TM_MODE
    await update.message.reply_text(
        "Режим: 🏷️ Товарные знаки.\n\n"
        "Отправьте название/ключевые слова — найду строки в Google Sheets и пришлю карточки.\n"
        "Команды:\n"
        "• /tm_reg — записи, где статус содержит «регистрация»\n"
        "• /tm_exp — записи, где статус содержит «экспертиза»",
        reply_markup=MAIN_KB
    )

async def tm_cmd_reg(update: Update, context: Any):
    await tm_process_search(update.effective_chat.id, ROW_MATCH_REGISTERED, context)

async def tm_cmd_exp(update: Update, context: Any):
    await tm_process_search(update.effective_chat.id, ROW_MATCH_EXPERTISE, context)

async def tm_handle_text(update: Update, context: Any):
    user_text = (update.message.text or "").strip()
    kws = re.split(r"\s+", user_text)
    await tm_process_search(update.effective_chat.id, lambda row: ROW_MATCH_KW(row, kws), context)

async def _work_handle_file(update: Update, context: Any, document: Document):
    raw = document.file_name or "file"
    fname = Path(raw).name
    base, ext = os.path.splitext(fname)
    ext = _normalize_ext(ext.lower().lstrip("."))
    if ext not in ("pdf", "docx", "doc"):
        await update.message.reply_text("Поддерживаются только PDF, DOCX, DOC для режима 'Работа с документами'.")
        return
    os.makedirs(DOC_FOLDER, exist_ok=True)
    ts = int(time.time())
    file_path = os.path.join(DOC_FOLDER, f"{base}_{ts}.{ext}")
    new_file = await context.bot.get_file(document.file_id)
    await new_file.download_to_drive(file_path)

    if ext == "pdf":
        key = _pdf_job_key(document)
        if _pdf_job_is_running(context, key):
            return
        sent = await update.message.reply_text("Принят PDF. Выполняю сжатие и OCR…")
        _pdf_job_start(context, key, sent.message_id)
        try:
            def _job(pdf_path: str):
                pdf_use = _compress_pdf_sync(pdf_path) if PDF_COMPRESS else pdf_path
                quick = extract_text_from_pdf(pdf_use) or ""
                if (not FORCE_OCR) and len(quick.strip()) > 50:
                    use_text = _normalize_confusables_ru(quick)
                else:
                    use_text = _ocr_pdf_to_text_sync(pdf_use, OCR_DPI, OCR_MAX_PAGES, POPPLER_PATH, TESS_LANG)
                use_text = postprocess(use_text)
                docx_path = _save_text_as_docx(use_text, f"{Path(fname).stem}_ocr")
                chunks = smart_split_text(use_text)
                return use_text, chunks, docx_path, pdf_use != pdf_path, pdf_use
            text, chunks, derived_docx, compressed, pdf_final = await asyncio.to_thread(_job, file_path)
            context.user_data.setdefault("work_docs", []).append({
                "name": fname, "path": file_path, "text": text or "", "chunks": chunks or [],
                "from_pdf": True, "docx_path": derived_docx, "pdf_processed": pdf_final,
            })
            try:
                await update.effective_chat.edit_message_text(
                    message_id=sent.message_id,
                    text=f"PDF добавлен в контекст. {'Сжатие выполнено. ' if PDF_COMPRESS else ''}"
                         f"OCR: {'успех' if (text or '').strip() else 'не удалось (использую извлечённый текст)'}."
                )
            except Exception:
                await update.message.reply_text(
                    f"PDF добавлен в контекст. {'Сжатие выполнено. ' if PDF_COMPRESS else ''}"
                    f"OCR: {'успех' if (text or '').strip() else 'не удалось (использую извлечённый текст)'}."
                )
        finally:
            _pdf_job_finish(context, key)
        return
    # DOC / DOCX
    if ext == "docx":
        text = extract_text_from_docx(file_path)
    else:
        text = extract_text_from_doc(file_path)

    if not (text or "").strip():
        await update.message.reply_text("Не удалось извлечь текст из документа. Убедитесь, что это не пустой/защищённый файл.")
        return
    chunks = smart_split_text(text)
    context.user_data.setdefault("work_docs", []).append({"name": fname, "path": file_path, "text": text, "chunks": chunks})
    await update.message.reply_text(f"Файл добавлен в контекст: {fname}. Документов в сессии: {len(context.user_data['work_docs'])}.")

async def handle_file(update: Update, context: Any):
    context.user_data.setdefault("work_docs", [])
    document: Document = update.message.document
    raw = document.file_name or "file"
    fname = Path(raw).name
    base, ext = os.path.splitext(fname)
    ext = _normalize_ext(ext.lower().lstrip("."))
    if ext not in ("pdf", "docx", "txt", "doc"):
        await update.message.reply_text("Поддерживаются только PDF, DOCX, DOC (и TXT для глобальной индексации).")
        return
    if context.user_data.get("mode") == "work":
        await _work_handle_file(update, context, document)
        return
    os.makedirs(DOC_FOLDER, exist_ok=True)
    ts = int(time.time())
    file_path = os.path.join(DOC_FOLDER, f"{base}_{ts}.{ext}")
    new_file = await context.bot.get_file(document.file_id)
    await new_file.download_to_drive(file_path)

    # Pre-calc hash for dedup
    try:
        file_hash = sha256_file(file_path)
    except Exception:
        file_hash = None

    manifest = load_manifest()
    if file_hash and file_hash in manifest.get("hashes", {}):
        await update.message.reply_text("Этот файл уже проиндексирован ранее. Можете задавать вопросы.")
        return
    # start indexing once; dedup by file_unique_id + hash
    key = _index_job_key(document.file_unique_id, file_hash)
    if _idx_is_running(context, key):
        return
    progress = await update.message.reply_text("Файл загружен. Индексация началась…")
    _idx_start(context, key, progress.message_id)

    try:
        def _index_job():
            try:
                added, total = index_file(file_path)
                return (True, added, total, None)
            except Exception as e:
                return (False, 0, 0, repr(e))
        ok, added, total, err = await asyncio.to_thread(_index_job)

        if ok:
            if added == 0:
                text = ("Файл загружен, но текст не извлечён. Возможно, это скан без текстового слоя.\n"
                        "Пришлите DOCX/TXT или PDF с текстом, либо включите OCR (tesseract+poppler/Docker).")
            else:
                if file_hash:
                    manifest.setdefault("hashes", {})[file_hash] = {"fname": os.path.basename(file_path), "time": int(time.time())}
                    save_manifest(manifest)
                text = f"Индексация завершена. Добавлено фрагментов: {added}. Всего: {total}. Теперь можно задавать вопросы."
        else:
            text = f"❌ Ошибка индексации: {err}"

        try:
            await update.effective_chat.edit_message_text(message_id=progress.message_id, text=text)
        except Exception:
            await update.message.reply_text(text)
    finally:
        _idx_finish(context, key)

async def _work_retrieve(docs: list, query: str, k: int = 6) -> list[str]:
    if not docs: return []
    all_chunks = [c for d in docs for c in d.get("chunks", [])]
    if not all_chunks: return []
    import numpy as np
    try:
        q = get_embedding(query)
    except Exception:
        return all_chunks[:k]
    embs = []
    for c in all_chunks:
        try: embs.append(get_embedding(c))
        except Exception: embs.append(None)
    scored = []
    for c, e in zip(all_chunks, embs):
        if e is None: continue
        denom = (np.linalg.norm(q) * np.linalg.norm(e)) or 1.0
        scored.append((float(q @ e / denom), c))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [c for _, c in scored[:k]] if scored else all_chunks[:k]

async def _reply_with_docx(update: Update, title: str, content: str, base_name: str):
    safe_base = re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", base_name)[:60] or "report"
    filename = f"{int(time.time())}_{safe_base}.docx"
    path = _save_docx_report(filename, title, content)
    with open(path, "rb") as f:
        await update.message.reply_document(InputFile(f, filename=filename))


def _get_last_pdf_doc(context) -> dict | None:
    """Возвращает последний загруженный документ из work_docs/general с расширением .pdf"""
    docs = (context.user_data.get("work_docs") or []) + (context.user_data.get("docs") or [])
    for rec in reversed(docs):
        name = rec.get("name") or ""
        if name.lower().endswith(".pdf"):
            return rec
    return None



async def pdf_to_txt(update: Update, context: Any):
    """Конвертирует последний загруженный PDF в TXT и отправляет файл пользователю."""
    rec = _get_last_pdf_doc(context)
    if not rec:
        await update.message.reply_text("Не найден загруженный PDF. Пришлите PDF-файл и повторите.")
        return
    text = (rec.get("text") or "").strip()
    if len(text) < 10:
        try:
            pdf_path = rec.get("path") or rec.get("pdf_processed") or ""
            if not pdf_path:
                await update.message.reply_text("Не удалось определить путь к PDF. Пришлите файл ещё раз.")
                return
            from app.services.extract import extract_text_from_pdf
            text = extract_text_from_pdf(pdf_path)
        except Exception as e:
            await update.message.reply_text(f"Ошибка извлечения текста: {e!r}")
            return
    try:
        text = postprocess(text or "")
    except Exception:
        pass
    if not (text or "").strip():
        await update.message.reply_text("Текст не распознан. Убедитесь, что PDF читабелен.")
        return
    safe_base = re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", (rec.get("name") or "document"))[:60]
    filename = f"{int(time.time())}_{safe_base}.txt"
    import tempfile, os
    with tempfile.NamedTemporaryFile(mode="w+", suffix=".txt", delete=False, encoding="utf-8") as tmp:
        tmp.write(text)
        tmp.flush()
        tmp_path = tmp.name
    try:
        with open(tmp_path, "rb") as f:
            await update.message.reply_document(InputFile(f, filename=filename), caption="Готово: PDF → TXT")
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass


async def doc_summary(update: Update, context: Any):
    docs = context.user_data.get("work_docs") or []
    if not docs:
        await update.message.reply_text("Нет документов в контексте. Пришлите DOC/DOCX/PDF и повторите.")
        return
    last = docs[-1]
    retrieved = await _work_retrieve([last], "краткое резюме документа с фокусом на тип работ/услуг/поставки/агентских/аренды", k=12)
    prompt = (
        "Ты — юридический ассистент. На основе КОНТЕКСТА составь краткое, но точное резюме договора.\n"
        "Обязательно:\n"
        "• Определи ТИП правоотношения и предмет: что именно делает исполнитель для заказчика (выбери из: выполнение работ, оказание услуг, поставка, агентские услуги, аренда), укажи формулировку из документа.\n"
        "• Стороны (наименования), срок, цена/порядок оплаты, ответственность/неустойки, расторжение, конфиденциальность/ПДн, права на результаты (если есть).\n"
        "• Избегай общих фраз «предмет договора» — пиши конкретное действие.\n"
        "Формат:\n"
        "— Тип/Предмет: …\n— Стороны: …\n— Срок: …\n— Цена/оплата: …\n— Ответственность: …\n— Расторжение: …\n— Конфиденциальность/ПДн: …\n— Права на РИД: …\n"
        "Используй только факты из контекста, без выдумок."
    )
    answer = generate_answer_with_gemini(prompt, retrieved) or ""
    if not answer.strip() or answer.strip().startswith("⚠️"):
        answer = _fallback_summary(last["text"])
    await send_long(update, f"📄 Резюме по: {last['name']}\n\n{answer}")
    await _reply_with_docx(update, f"Резюме: {last['name']}", answer, f"summary_{Path(last['name']).stem}")

async def doc_check(update: Update, context: Any):
    docs = context.user_data.get("work_docs") or []
    if not docs:
        await update.message.reply_text("Нет документов в контексте. Сначала загрузите файл.")
        return
    last = docs[-1]
    retrieved = await _work_retrieve([last], "проверка договора на ошибки/риски/пробелы/несоответствия", k=16)
    prompt = (
        "Ты — опытный договорной юрист. Проверь договор на ошибки и риски строго по КОНТЕКСТУ.\n"
        "Укажи: спорные формулировки, односторонние условия, пробелы (не хватает условий), конфликты разделов,\n"
        "риски по цене/срокам/ответственности/расторжению/РИД/конфиденциальности/ПДн и дай конкретные правки (как переписать пункты).\n"
        "Структура:\n1) Краткий итог рисков (3–6 пунктов)\n2) Детальный разбор по разделам с цитатами\n3) Список предложенных правок (bullet list)"
    )
    answer = generate_answer_with_gemini(prompt, retrieved) or ""
    if not answer.strip() or answer.strip().startswith("⚠️"):
        answer = _fallback_check(last["text"])
    await send_long(update, f"🔎 Проверка договора: {last['name']}\n\n{answer}")
    await _reply_with_docx(update, f"Проверка договора: {last['name']}", answer, f"check_{Path(last['name']).stem}")




async def doc_compare(update, context):
    docs = context.user_data.get("work_docs") or []
    if len(docs) < 2:
        await update.message.reply_text("Нужно минимум два документа. Загрузите ещё один и повторите /doc_compare.")
        return
    a, b = docs[-2], docs[-1]
    text_a = a.get("text") or ""
    text_b = b.get("text") or ""

    raw_pairs = _pair_sentences(text_a, text_b)
    logging.getLogger("semantic-bot").info("[Gemini] call site: got %s raw pairs before filtering", len(raw_pairs))
    try:
        filtered = await _gemini_semantic_filter_sentence_pairs(raw_pairs, api_key=None)
    except Exception as e:
        logging.getLogger("semantic-bot").exception("[Gemini] filtering call failed — using heuristic: %r", e)
        filtered = [(o, r) for (o, r) in raw_pairs if _heuristic_semantic_diff(o, r)]

    filename = f"{int(time.time())}_changes_{Path(a.get('name','original')).stem}_to_{Path(b.get('name','recognized')).stem}.docx"
    path = _save_compare_report_with_adapter(
        filtered_pairs=filtered,
        filename=filename,
        a_name=a.get('name','original.docx'),
        b_name=b.get('name','recognized.docx'),
        text_a=text_a,
        text_b=text_b,
    )

    try:
        with open(path, "rb") as f:
            await update.message.reply_document(InputFile(f, filename=os.path.basename(path)))
    except Exception:
        await update.message.reply_text(f"Файл подготовлен: {path}")


async def doc_clear(update: Update, context: Any):
    docs = context.user_data.get("work_docs") or []
    deleted_count = 0
    for d in docs:
        for key in ("path", "docx_path", "pdf_processed"):
            p = d.get(key)
            if p and os.path.exists(p):
                try:
                    os.remove(p); deleted_count += 1
                except Exception:
                    pass
    context.user_data["work_docs"] = []
    await update.message.reply_text(f"Очищено файлов: {deleted_count}. Контекст теперь пуст.")

async def handle_text(update: Update, context: Any):
    user_query = (update.message.text or "").strip()
    mode = context.user_data.get("mode", "docs")

    if user_query == AI_LABEL:
        await ai_mode(update, context); return
    if user_query == DOCS_LABEL:
        await docs_mode(update, context); return
    if user_query == TM_LABEL:
        await tm_mode(update, context); return
    if user_query == WORK_LABEL:
        await work_mode(update, context); return

    if mode == "ai":
        uid = update.effective_user.id
        if not is_admin(uid):
            used = get_usage(uid)
            if used >= DAILY_FREE_LIMIT:
                await update.message.reply_text(
                    "Достигнут дневной лимит бесплатных обращений к ИИ (10). Возвращайтесь завтра или обратитесь к администратору.",
                    reply_markup=MAIN_KB,
                )
                return
        def _ai_job():
            try:
                return generate_direct_ai_answer(user_query)
            except Exception as e:
                return f"⚠️ Ошибка при обработке запроса: {repr(e)}"
        answer = await asyncio.to_thread(_ai_job)
        if answer and not is_admin(uid):
            inc_usage(uid)
        await send_long(update, answer)
        if not is_admin(uid):
            left = max(0, DAILY_FREE_LIMIT - get_usage(uid))
            await update.message.reply_text(f"Остаток на сегодня: {left}.")
        return
    if mode == "work":
        docs = context.user_data.get("work_docs") or []
        if not docs:
            await update.message.reply_text("Пока нет документов в контексте. Пришлите DOC/DOCX/PDF.")
            return
        retrieved = await _work_retrieve(docs, user_query, k=RETRIEVAL_K)
        answer = generate_answer_with_gemini(user_query, retrieved) or ""
        if not answer.strip() or answer.strip().startswith("⚠️"):
            answer = "Краткая справка по загруженным документам:\n" + _fallback_summary("\n\n".join(retrieved) if retrieved else docs[-1]["text"])
        await send_long(update, answer)
        return
    if mode == "tm":
        low = user_query.lower()
        if low.startswith("/tm_reg"):
            await tm_cmd_reg(update, context); return
        if low.startswith("/tm_exp"):
            await tm_cmd_exp(update, context); return
        await tm_handle_text(update, context)
        return
    if not (os.path.exists(INDEX_FILE) and os.path.exists(TEXTS_FILE)):
        await update.message.reply_text("Нет проиндексированных документов. Сначала загрузите файл.")
        return
    def _answer_job():
        try:
            chunks = retrieve_chunks(user_query, k=RETRIEVAL_K)
            return generate_answer_with_gemini(user_query, chunks)
        except Exception as e:
            return f"⚠️ Ошибка при обработке запроса: {repr(e)}"

    answer = await asyncio.to_thread(_answer_job)
    if not answer:
        await update.message.reply_text("⚠️ Пустой ответ.")
    else:
        await send_long(update, answer)

async def error_handler(update: object, context: Any) -> None:
    logger.exception("Unhandled error while processing update: %s", update)

from telegram import InputFile
from app.config import INDEX_FILE, TEXTS_FILE, DOCMETA_FILE

async def download_index(update: Update, context: Any):
    # Только для админов
    if not is_admin(update.effective_user.id):
        await update.message.reply_text("Команда доступна только администраторам.")
        return
    try:
        path = INDEX_FILE
        await _send_file_safely(update, path, "index.faiss", "FAISS-индекс")
    except Exception as e:
        await update.message.reply_text(f"Ошибка: {e}")


async def download_texts(update: Update, context: Any):
    # Только для админов
    if not is_admin(update.effective_user.id):
        await update.message.reply_text("Команда доступна только администраторам.")
        return
    try:
        path = TEXTS_FILE
        await _send_file_safely(update, path, "texts.pkl", "Чанки текстов")
    except Exception as e:
        await update.message.reply_text(f"Ошибка: {e}")


async def download_meta(update: Update, context: Any):
    # Только для админов
    if not is_admin(update.effective_user.id):
        await update.message.reply_text("Команда доступна только администраторам.")
        return
    try:
        path = DOCMETA_FILE
        await _send_file_safely(update, path, "docmeta.json", "Метаданные документов")
    except Exception as e:
        await update.message.reply_text(f"Ошибка: {e}")


async def _count_tokens_safe(model, text: str) -> int:
    try:
        ct = await asyncio.to_thread(model.count_tokens, text)
        for attr in ("total_tokens", "total_token_count", "token_count"):
            if hasattr(ct, attr):
                val = getattr(ct, attr)
                try:
                    return int(getattr(val, "value", val))
                except Exception:
                    try:
                        return int(val)
                    except Exception:
                        pass
        return 0
    except Exception:
        return (len(text) + 3) // 4

def _batch_pairs_by_limits(items: List[Tuple[str, str]], max_pairs: int = 60, max_chars: int = 18000):
    batch, cur_chars = [], 0
    for (o, r) in items:
        o_t, r_t = _truncate(o), _truncate(r)
        payload = json.dumps({"original": o_t, "recognized": r_t}, ensure_ascii=False)
        if (len(batch) >= max_pairs) or (cur_chars + len(payload) > max_chars):
            if batch:
                yield batch
            batch, cur_chars = [], 0
        batch.append((o_t, r_t))
        cur_chars += len(payload)
    if batch:
        yield batch

def _build_prompt(items: List[Tuple[str, str]]) -> str:
    return (
        """Ты — юридический редактор. Даны пары предложений из договора.
- ORIGINAL — из исходного DOCX
- RECOGNIZED — из распознанного PDF (OCR-ошибки возможны)

Нужно оставить ТОЛЬКО пары со смысловым отличием (разные значения).
Игнорируй различия в пробелах, регистре и пунктуации.
Если в паре различаются ЧИСЛА или даты — это смысловое отличие.

Верни ЧИСТЫЙ JSON-массив объектов (без подсказок/комментариев, БЕЗ Markdown-кода):
[
  {"original": "...", "recognized": "..."},
  ...
]

PAIRS:
"""
        + json.dumps([{"original": o, "recognized": r} for (o, r) in items], ensure_ascii=False)
    )


async def _gemini_semantic_filter_sentence_pairs(pairs, api_key: str | None):
    log = logging.getLogger("semantic-bot")
    log.info("[Gemini] entry: pairs=%s", len(pairs) if pairs else 0)

    if not pairs:
        log.info("[Gemini] no pairs — nothing to filter")
        return []

    key = api_key or os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
    if not key:
        log.info("[Gemini] API key missing — skip filtering; using heuristic")
        return [(o, r) for (o, r) in pairs if _heuristic_semantic_diff(o, r)]

    try:
        import google.generativeai as genai
        genai.configure(api_key=key)
        model = genai.GenerativeModel("gemini-1.5-flash")
    except Exception as e:
        log.exception("[Gemini] init failed: %r — using heuristic", e)
        return [(o, r) for (o, r) in pairs if _heuristic_semantic_diff(o, r)]

    # smaller batches
    def _batch_pairs_by_limits(items, max_pairs: int = 40, max_chars: int = 12000):
        batch, cur_chars = [], 0
        for (o, r) in items:
            o_t, r_t = (o or "").strip(), (r or "").strip()
            payload = json.dumps({"original": o_t, "recognized": r_t}, ensure_ascii=False)
            if (len(batch) >= max_pairs) or (cur_chars + len(payload) > max_chars):
                if batch:
                    yield batch
                batch, cur_chars = [], 0
            batch.append((o_t, r_t))
            cur_chars += len(payload)
        if batch:
            yield batch

    batches = list(_batch_pairs_by_limits(pairs, max_pairs=40, max_chars=12000))
    log.info("[Gemini] filtering enabled, total pairs: %s, batches: %s", len(pairs), len(batches))

    filtered_total = []
    disable_gemini_for_run = False

    for i, batch in enumerate(batches):
        if disable_gemini_for_run:
            filtered_total.extend([(o, r) for (o, r) in batch if _heuristic_semantic_diff(o, r)])
            continue

        prompt = _build_prompt(batch)
        try:
            tok = await _count_tokens_safe(model, prompt)
            if tok > 30000 and len(batch) > 1:
                mid = len(batch) // 2
                batches[i:i+1] = [batch[:mid], batch[mid:]]
                log.info("[Gemini] split batch %s due to token estimate %s", i+1, tok)
                continue
        except Exception:
            pass

        attempt, last_err = 0, None
        while attempt < 2 and not disable_gemini_for_run:
            attempt += 1
            try:
                log.info("[Gemini] batch %s/%s, size: %s (attempt %s)", i+1, len(batches), len(batch), attempt)
                resp = await asyncio.to_thread(
                    lambda: model.generate_content(
                        contents=[{"role": "user", "parts": [{"text": prompt}]}],
                        generation_config={
                            "response_mime_type": "application/json",
                            "temperature": 0.0,
                            "candidate_count": 1,
                        },
                    )
                )
                raw = getattr(resp, "text", None) or getattr(resp, "output_text", None) or ""
                data = _extract_json_array(raw)
                if not isinstance(data, list):
                    raise ValueError("Gemini did not return a JSON array")

                kept = 0
                for item in data:
                    o = (item.get("original") or "").strip() if isinstance(item, dict) else ""
                    r = (item.get("recognized") or "").strip() if isinstance(item, dict) else ""
                    if o and r and _heuristic_semantic_diff(o, r):
                        filtered_total.append((o, r)); kept += 1
                log.info("[Gemini] batch %s kept: %s/%s", i+1, kept, len(batch))
                break
            except Exception as e:
                last_err = e
                msg = f"{e.__class__.__name__}: {e}"
                log.warning("[Gemini] batch %s failed: %s", i+1, msg)
                if ("ResourceExhausted" in e.__class__.__name__) or ("quota" in str(e).lower()):
                    log.error("[Gemini] quota hit — disabling Gemini for this run, using heuristic for remaining batches")
                    disable_gemini_for_run = True
                    break
                await asyncio.sleep(0.6 * attempt)

        if last_err and (attempt >= 2 or disable_gemini_for_run):
            log.error("[Gemini] batch %s fallback to heuristic due to persistent error: %r", i+1, last_err)
            filtered_total.extend([(o, r) for (o, r) in batch if _heuristic_semantic_diff(o, r)])

    log.info("[Gemini] filtered pairs total: %s (from %s)", len(filtered_total), len(pairs))
    return filtered_total

def _split_sentences(text: str):
    text = (text or "").strip()
    if not text:
        return []
    parts = [p.strip() for p in _SENT_SPLIT_RE.split(text) if p.strip()]
    return parts

def _pair_sentences(a_text: str, b_text: str):
    a_s = _split_sentences(a_text)
    b_s = _split_sentences(b_text)
    sm = difflib.SequenceMatcher(a=a_s, b=b_s, autojunk=False)
    pairs = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        if tag == "replace":
            for i in range(max(i2-i1, j2-j1)):
                o = a_s[i1+i] if i1+i < i2 else ""
                r = b_s[j1+i] if j1+i < j2 else ""
                if o or r:
                    pairs.append((o, r))
        elif tag == "delete":
            for i in range(i1, i2):
                pairs.append((a_s[i], ""))
        elif tag == "insert":
            for j in range(j1, j2):
                pairs.append(("", b_s[j]))
    return pairs

def _save_docx_semantic_pairs(filename: str, doc_name_a: str, doc_name_b: str, pairs):
    os.makedirs(DOC_FOLDER, exist_ok=True)
    path = os.path.join(DOC_FOLDER, filename)
    doc = DocxDocument()
    doc.add_heading(f"Смысловые отличия: {doc_name_a} → {doc_name_b}", level=1)
    if not pairs:
        doc.add_paragraph("Смысловых отличий не найдено.")
        doc.save(path); return path

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Оригинал (DOCX)"
    hdr_cells[1].text = "Распознанный (OCR)"
    for (o, r) in pairs:
        row = table.add_row().cells
        row[0].text = o or "—"
        row[1].text = r or "—"
    doc.save(path)
    return path



def _extract_json_array(raw: str):
    if not raw:
        return []
    s = raw.strip()
    if s.startswith("```"):
        s = s.strip("`").strip()
        if s.lower().startswith("json"):
            s = s[4:].strip()
    try:
        data = json.loads(s)
        return data if isinstance(data, list) else []
    except Exception:
        pass
    l = s.find("["); r = s.rfind("]")
    if l != -1 and r != -1 and r > l:
        chunk = s[l:r+1].strip()
        try:
            data = json.loads(chunk)
            return data if isinstance(data, list) else []
        except Exception:
            parts = re.split(r"\]\s*[\r\n]+\s*\[", chunk.strip()[1:-1])
            items = []
            for part in parts:
                pj = "[" + part.strip() + "]"
                try:
                    items.extend(json.loads(pj))
                except Exception:
                    pass
            if items:
                return items
    objs = re.findall(r"\{[^{}]*\}", s, flags=re.DOTALL)
    if objs:
        try:
            return [json.loads(o) for o in objs]
        except Exception:
            pass
    return []





def _save_compare_report_with_adapter(filtered_pairs, filename, a_name, b_name, text_a, text_b):
    """
    Universal caller for project _save_docx_compare_tables with signature detection.
    Tries keyword invocation via inspect.signature; then tries several positional orders.
    Fallbacks to DOCX/TSV writer.
    """
    if "_save_docx_compare_tables" in globals():
        fn = _save_docx_compare_tables
        # Keyword attempt
        try:
            sig = inspect.signature(fn)
            params = list(sig.parameters.keys())
            kwargs = {}
            mapping = {
                "pairs": filtered_pairs,
                "filtered_pairs": filtered_pairs,
                "filename": filename,
                "doc_name_a": a_name,
                "doc_name_b": b_name,
                "text_a": text_a,
                "text_b": text_b,
            }
            for p in params:
                if p in mapping:
                    kwargs[p] = mapping[p]
            if kwargs and ("filename" in kwargs and ("pairs" in kwargs or "filtered_pairs" in kwargs)):
                return fn(**kwargs)
        except Exception:
            logging.getLogger("semantic-bot").exception("Project _save_docx_compare_tables kwargs call failed; try positional")

        # Positional attempts
        orders = [
            (filename, filtered_pairs, b_name, text_a, text_b),
            (filename, filtered_pairs, a_name, b_name, text_a, text_b),
            (filtered_pairs, filename, b_name, text_a, text_b),
            (filtered_pairs, filename, a_name, b_name, text_a, text_b),
        ]
        for args in orders:
            try:
                return fn(*args)
            except TypeError:
                continue
            except Exception:
                logging.getLogger("semantic-bot").exception("Project _save_docx_compare_tables positional attempt failed; trying next")
        logging.getLogger("semantic-bot").error("Project _save_docx_compare_tables signature mismatch — using fallback writer")

    # Fallback writer
    try:
        from docx import Document as DocxDocument
        os.makedirs("generated_reports", exist_ok=True)
        path = os.path.join("generated_reports", filename)
        doc = DocxDocument()
        doc.add_heading(f"Смысловые отличия: {a_name} → {b_name}", level=1)
        if not filtered_pairs:
            doc.add_paragraph("Смысловых отличий не найдено.")
        else:
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Оригинал (DOCX)"
            table.rows[0].cells[1].text = "Распознанный (OCR)"
            for o, r in filtered_pairs:
                row = table.add_row().cells
                row[0].text = o or "—"
                row[1].text = r or "—"
        doc.save(path)
        return path
    except Exception:
        os.makedirs("generated_reports", exist_ok=True)
        path = os.path.join("generated_reports", filename.replace(".docx", ".tsv"))
        with open(path, "w", encoding="utf-8") as f:
            f.write("Оригинал\tРаспознанный\n")
            for o, r in filtered_pairs:
                line_o = (o or "").replace("\t", " ")
                line_r = (r or "").replace("\t", " ")
                f.write(line_o + "\t" + line_r + "\n")
        return path


from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters
from app.config import TELEGRAM_BOT_TOKEN

def build_application():
    app = ApplicationBuilder().token(TELEGRAM_BOT_TOKEN).build()

    # === команды ===
    try:
        app.add_handler(CommandHandler("start", start))
    except Exception:
        pass
    try:
        app.add_handler(CommandHandler("help", help_command))
    except Exception:
        pass
    try:
        app.add_handler(CommandHandler("download_index", download_index))
        app.add_handler(CommandHandler("download_texts", download_texts))
        app.add_handler(CommandHandler("download_meta", download_meta))
    except Exception:
        pass

    try:
        app.add_handler(CommandHandler("tm_reg", tm_cmd_reg))
        app.add_handler(CommandHandler("tm_exp", tm_cmd_exp))
    except Exception:
       pass

# === переключатели режимов ===
    try:
        app.add_handler(MessageHandler(filters.TEXT & filters.Regex(f"^{re.escape(TM_LABEL)}$"), tm_mode))
        app.add_handler(MessageHandler(filters.TEXT & filters.Regex(f"^{re.escape(WORK_LABEL)}$"), work_mode))
        app.add_handler(MessageHandler(filters.TEXT & filters.Regex(f"^{re.escape(DOCS_LABEL)}$"), docs_mode))
        app.add_handler(MessageHandler(filters.TEXT & filters.Regex(f"^{re.escape(AI_LABEL)}$"), ai_mode))
    except Exception:
        pass

    # === обработчики входящих сообщений и файлов ===
    try:
        app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
        app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    except Exception:
        pass

    # === обработчик ошибок ===
    try:
        app.add_error_handler(error_handler)
    except Exception:
        pass

    return app

__all__ = ["build_application"]
